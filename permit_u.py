#-*- coding: utf-8 -*-
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlrd
import time
todolist=[]
workbook = xlrd.open_workbook('todolist.xlsx')
table = workbook.sheets()[0]
nrows = table.nrows
ncols = table.ncols
for i in range(0,nrows):
	row_list=[]
	rowValues= table.row_values(i) 
	todolist.append(rowValues)
todolist.remove(todolist[0])
#print (todolist)
for row in todolist:
	print (row)
	ymd_permit=row[7].split('.')
	if int(ymd_permit[1])==1:
		ymd_todo=ymd_permit[0]+'年'+'12'+'月'+ymd_permit[-1]+'日'
	else:
		ymd_todo=str(int(ymd_permit[0])+1)+'年'+str(int(ymd_permit[1])-1)+'月'+ymd_permit[-1]+'日'
	ymd_permit=ymd_permit[0]+'年'+ymd_permit[1]+'月'+ymd_permit[-1]+'日'
	#print (ymd_todo)
	#doc = Document()
	doc = Document('.\\template_blank.docx')
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = docx.shared.Pt(16)

	#doc.add_heading('Document Title', 0)
	title=doc.add_paragraph()
	title_run=title.add_run('证  明\n\n')
	font = title_run.font
	#font.name = 'Calibri'
	font.bold = True
	font.size = docx.shared.Pt(22)
	paragraph_format = title.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	date_txt='大理州公安局出入境管理支队:'
	date=doc.add_paragraph()
	date_run=date.add_run(date_txt)
	font = date_run.font
	font.size = docx.shared.Pt(18)
	if '年'in str(row[8]):
		JW202=row[8]
	else:
		ymd_202=str(row[8]).split('.')
		ymd_202=ymd_202[0]+'年'+ymd_202[1]+'月'
		JW202=ymd_202

	body=doc.add_paragraph()
	
	body1='    兹有我校'
	body1_run=body.add_run(body1)
	paragraph_format = body.paragraph_format
	paragraph_format.line_spacing = docx.shared.Pt(30)
	fill1='%s'%row[5]
	fill1_run=body.add_run(fill1)
	font = fill1_run.font
	font.underline = True
	
	body2='籍学生'
	body2_run=body.add_run(body2)
	fill2='%s'%row[2].upper()
	fill2_run=body.add_run(fill2)
	font = fill2_run.font
	font.underline = True
	
	body3=', 性别：'
	body3_run=body.add_run(body3)
	fill3='%s'%row[4]
	fill3_run=body.add_run(fill3)
	font = fill3_run.font
	font.underline = True

	body4='，护照号码为'
	body4_run=body.add_run(body4)
	fill4='%s'%row[6]
	fill4_run=body.add_run(fill4)
	font = fill4_run.font
	font.underline = True

	body5='。其居留许可到期，现需办理学习居留许可延期手续，时间为'
	body5_run=body.add_run(body5)
	fill5='%s'%ymd_todo
	fill5_run=body.add_run(fill5)
	font = fill5_run.font
	font.underline = True

	body6='，请按有关规定给予办理为谢。\n    居留许可到期时间：'
	body6_run=body.add_run(body6)
	fill6='%s'%ymd_permit
	fill6_run=body.add_run(fill6)
	font = fill6_run.font
	font.underline = True

	body7='\n    JW202表到期时间：'
	body7_run=body.add_run(body7)
	fill7='%s'%JW202
	fill7_run=body.add_run(fill7)
	font = fill7_run.font
	font.underline = True

	body8='\n\n\n\n'
	body8_run=body.add_run(body8)
	
	year='年'
	month='月'
	day='日'
	cc=time.localtime(time.time())
	end_txt='大理大学留学生教育服务中心\n%s'%str(cc.tm_year)+year+str(cc.tm_mon)+month+str(cc.tm_mday)+day
	end=doc.add_paragraph()
	end_run=end.add_run(end_txt)
	paragraph_format = end.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	
	doc.save('%s.docx'%row[1])
print ('已成功生成所有文件！')
time.sleep(7)