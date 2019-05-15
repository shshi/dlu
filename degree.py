#-*- coding: utf-8 -*-
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlrd
import time
import datetime

workbook = xlrd.open_workbook('xxx.xlsx')
sheet_names= workbook.sheet_names()
sheet1 = workbook.sheet_by_name(sheet_names[0])
name_list = sheet1.col_values(4)
for i in name_list:
	i=i.strip()

name=input('请输入学生英文名：')
name=name.upper()

if name in name_list:
	row_num=name_list.index(name)
	rowValues= sheet1.row_values(row_num)
	print (rowValues)
else:
	print ('数据库中没有匹配到该生信息')

if '男' in rowValues[7]:
	title='Mr. '
	sex='his'
elif '女' in rowValues[7]:
	title='Ms. '
	sex='her'
name=title+rowValues[4]
doc = Document('template_blank.docx')
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = docx.shared.Pt(12)

if '印度' in rowValues[6]:
	body_ind='To: Consulate General of India,\nGuangzhou,\nChina\n\nSubject:\nAuthenticity Confirmation on Certificates of %s\n'%name
	head=doc.add_paragraph()
	head_run=head.add_run(body_ind)
	by='MOE of China and Medical Council of India'
elif '孟加拉' in rowValues[6]:
	body_bgd='To: BANGLADESH MEDICAL & DENTAL COUNCIL,\nDhaka,\nBangladesh.\n\nSubject:\nAuthenticity Confirmation on Certificates of %s\n'%name
	head=doc.add_paragraph()
	head_run=head.add_run(body_bgd)
	by='MOE of China and BANGLADESH MEDICAL & DENTAL COUNCIL'
else:
	print ('此学生国籍暂无模板，使用印度籍模板代替！')
	body_ind='To: Consulate General of India,\nGuangzhou,\nChina\n\nSubject:\nAuthenticity Confirmation on Certificates of %s\n'%name
	head=doc.add_paragraph()
	head_run=head.add_run(body_ind)
	by='MOE of China and Medical Council of India'

today = datetime.date.today()
date_txt='Date: {:%b %d, %Y}\n'.format(today)
#print (date_txt)
date=doc.add_paragraph()
date_run=date.add_run(date_txt)
paragraph_format = date.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

body3='''Dear Consul: \n\nAfter University's verification, the certificates of %s were issued by Dali University. These documents are authentic and effective documents.\n\nHere are replies for other questions: \n\n1. %s studied at Dali University from November 2011 to November 2018, majored in Bachelor of Medicine and Bachelor of Surgery.\n\n2. %s has been studying in Dali University without any transferring to other universities since %s enrollment into Dali University.\n\n3. Certificates issued by Dali University are recognized by %s. Student’s qualification can be recognized for registration as medical practitioner in China only after %s passing relevant examination.\n\nYours faithfully,\n\n\nLiu Fengqin\nDeputy Director\nInternational Students Administration, Education & Service Center for International Students, Dali University\nNo.2 Hongsheng Road, Dali,Yunnan, P.R.CHINA, 671003\nEmail: dalijoe2005@aliyun.com\nTelephone: +86-(0)872-221-8978      Fax: +86-(0)872-221-8979\nhttp://www.dali.edu.cn'''%(name, name, name, sex, by, sex)
main=doc.add_paragraph()
main_run=main.add_run(body3)

doc.save('%s.docx'%rowValues[4])

print ('学位证明已生成！')
time.sleep(7)
