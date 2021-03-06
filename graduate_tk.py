#-*- coding: utf-8 -*-
import os
import time
import datetime
import docx
import xlrd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter.messagebox import *

class Graduate:
	def __init__(self,root):
		self.todolist=[]
		workbook = xlrd.open_workbook('graduate.xlsx')
		table = workbook.sheets()[0]
		nrows = table.nrows
		ncols = table.ncols
		for i in range(0,nrows):
			row_list=[]
			rowValues= table.row_values(i) 
			self.todolist.append(rowValues)
		self.todolist.remove(self.todolist[0])
		
		frame=tk.Frame(root)
		frame.pack()
		tk.Button(frame,text='No Criminal Certificate',command=self.loopNCC).pack(side=tk.LEFT)
		tk.Button(frame,text='Migration Certificate',command=self.loopMC).pack(side=tk.LEFT)
		tk.Button(frame,text='Study Certificate (Chinese)',command=self.loopSC_CN).pack(side=tk.LEFT)
		tk.Button(frame,text='Study Certificate (English)',command=self.loopSC_EN).pack(side=tk.LEFT)
		frame=tk.Frame(root)
		frame.pack()

		tk.Label(root,text='\nPowered by@ShaoTech\nCopyriht © 大理大学留学生教育服务中心 版权所有\n如有问题请联系：石少华，Email: shi.sh@foxmail.com').pack(side=tk.BOTTOM)
		README_btn=tk.Button(frame,text='使用说明',command=self.ReadMe)
		README_btn.pack(side=tk.TOP)

	def ReadMe(self):
		#os.system('README.txt')
		showinfo('使用说明','\n此程序所在文件夹中需有如下两个文件：\n1. 命名为graduate.xlsx的Excel文件。此文件为待做学生信息文件，Excel第一行为表头，从左到右应依次为：\n序号，学号，护照姓名，中文名，性别(格式为Male/Female)，护照号，生日(英文)，国籍(英文)\n2. 命名为template_blank.docx的大理大学留学生教育服务中心模版文件。若无此文件，生成的文件中将无页眉页脚。\n\n使用方法：\n1. 将待做学生信息写入graduate.xlsx\n2. 直接双击运行graduate_tk.py')
		
	def NCC(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('CERTIFICATE OF NO CRIMINAL RECORD')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		date=doc.add_paragraph()
		date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		now = datetime.datetime.now()
		today=now.strftime("%B %d, %Y")
		date_txt='\n%s\n'%today
		date_txt='\nJune 25'
		date.add_run(date_txt)
		date.add_run('th').font.superscript = True
		date.add_run(', 2019\n')
		
		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex:
			Sex='Ms. '
			HeShe='her'
		else:
			Sex='Mr. '
			HeShe='his'
		Name=row[2]
		PPNo=row[5]
		Grade='20'+row[1][2:4]
		if row[1][0]=='Y' or row[1][0]=='y':
			Grade=row[1][1:5]
		body_txt='''This is to certify that %s%s (passport No. %s; student No. %s) has no disciplinary records against the rules and regulations of Dali University and has no Chinese judicial office records of committing any offense against Chinese criminal laws during %s study in P.R. China from September %s to till now.\n\nCertified by\n'''%(Sex,Name,PPNo,RollNo,HeShe,Grade)

		body=doc.add_paragraph()
		body_run=body.add_run(body_txt)
		body.paragraph_format.line_spacing = docx.shared.Pt(30)

		inscription='\n\n\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003, P. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8978   Fax:+86-872-221-8979'	
		ins=doc.add_paragraph()
		ins.paragraph_format.line_spacing = docx.shared.Pt(25)
		ins_run=ins.add_run(inscription)
		ins_run.font.bold = True

		today_appen = datetime.date.today()

		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
			
		else:
			doc.save('NCC_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')
			
	def MC(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('MIGRATION CERTIFICATE')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		date=doc.add_paragraph()
		date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		now = datetime.datetime.now()
		today=now.strftime("%B %d, %Y")
		date_txt='\n%s\n'%today
		date_txt='\nJune 25'
		date.add_run(date_txt)
		date.add_run('th').font.superscript = True
		date.add_run(', 2019\n')

		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex:
			Sex='Ms. '
			HeShe='her'
		else:
			Sex='Mr. '
			HeShe='his'
		Name=row[2]
		PPNo=row[5]
		DOB=row[6]
		Nation=row[7]

		body_txt='''This university has no objection to the admission of %s%s, from %s, bearing Dali University’s registration No. %s, passport No. %s, born on %s, for %s further study in any institution or university in any country.\n\nWe wish %s success in life.\n\n\n'''%(Sex,Name,Nation,RollNo,PPNo,DOB,HeShe,HeShe)
		infoNeed=list([Sex,Name,Nation,RollNo,PPNo,DOB,HeShe,HeShe])
		if len([a for a in infoNeed if a.strip()==''])>0:
			print ('有缺失信息，请补全后再试！')
			showinfo('提示','有必填信息缺失，请补全后再试！')
		body=doc.add_paragraph()
		body_run=body.add_run(body_txt)
		paragraph_format = body.paragraph_format
		paragraph_format.line_spacing = docx.shared.Pt(30)
		inscription='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003, P. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8978   Fax:+86-872-221-8979'
		ins_run=body.add_run(inscription)
		font = ins_run.font
		font.bold = True

		today_appen = datetime.date.today()

		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
			
		else:
			doc.save('MC_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')

	def SC_CN(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('在 读 证 明\n\n')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		
		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex:
			Sex='女'
		else:
			Sex='男'
		Name=row[2]
		Name_CN=row[3]
		PPNo=row[5]
		Nation=row[7]
		if 'India' in Nation:
			Nation='印度'
		elif 'Nepal' in Nation:
			Nation='尼泊尔'
		elif 'Pakistan' in Nation:
			Nation='巴基斯坦'
		elif 'Bangladesh' in Nation:
			Nation='孟加拉国'
		elif 'Ivory Cost' in Nation:
			Nation='科特迪瓦'
		elif 'Laos' in Nation:
			Nation='老挝'
		elif 'Cambodia' in Nation:
			Nation='柬埔寨'
		elif 'Tanzania' in Nation:
			Nation='坦桑尼亚'
		elif 'Viet' in Nation:
			Nation='越南'
		elif 'Somalia' in Nation:
			Nation='索马里'
		elif 'Burma' in Nation or 'Myanmar' in Nation:
			Nation='缅甸'
		elif 'Zambia' in Nation:
			Nation='赞比亚'
		elif 'Yemen' in Nation:
			Nation='也门'
		elif 'Mongolia' in Nation:
			Nation='蒙古'

		body=doc.add_paragraph()
		body.paragraph_format.line_spacing = docx.shared.Pt(30)
		
		body.add_run('    兹证明')
		body.add_run('%s'%Name).font.underline = True

		body.add_run('，性别')
		body.add_run('%s'%Sex).font.underline = True
		
		body.add_run('，中文名')
		body.add_run('%s'%Name_CN).font.underline = True

		body.add_run('，国籍')
		body.add_run('%s'%Nation).font.underline = True

		body.add_run('，护照号码')
		body.add_run('%s'%PPNo).font.underline = True

		body.add_run('，为我校')
		body.add_run('临床医学院2013级临床医学专业本科生').font.underline = True
		body.add_run('。该生于2013年10月入学，2019年7月从我校毕业。\n\n\t特此证明。\n\n')

		infoNeed=list([Sex,Name,Name_CN,PPNo,Nation])
		if len([a for a in infoNeed if a.strip()==''])>0:
			print ('有缺失信息，请补全后再试！')
			showinfo('提示','有必填信息缺失，请补全后再试！')

		year='年'
		month='月'
		day='日'
		cc=time.localtime(time.time())
		end_txt='大理大学留学生教育服务中心\n%s'%str(cc.tm_year)+year+str(cc.tm_mon)+month+str(cc.tm_mday)+day
		#end_txt='大理大学留学生教育服务中心\n2019年6月25日'
		end=doc.add_paragraph()
		end_run=end.add_run(end_txt)
		paragraph_format = end.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		today_appen = datetime.date.today()
		
		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
			
		else:
			doc.save('SC_CN_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')

	def SC_EN(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('STUDY CERTIFICATE')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		date=doc.add_paragraph()
		date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		now = datetime.datetime.now()
		today=now.strftime("%B %d, %Y")
		date_txt='\n%s\n'%today
		#date_txt='\nJune 25'
		date.add_run(date_txt)
		#date.add_run('th').font.superscript = True
		#date.add_run(', 2019\n')

		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex:
			Sex='Ms. '
			HeShe='her'
			He='She'
		else:
			Sex='Mr. '
			HeShe='his'
			He='He'
		Name=row[2]
		PPNo=row[5]
		DOB=row[6]
		Nation=row[7]
		Major='Clinical Medicine'
		School='Clinical Medicine College'
		Grade='20'+row[1][2:4]
		if row[1][0]=='Y' or row[1][0]=='y':
			Grade=row[1][1:5]
		toyear=now.strftime("%Y")
		year_count=int(toyear)-int(Grade)
		if year_count ==0:
			year_count='first'
		elif year_count ==1:
			year_count='first'
		elif year_count ==2:
			year_count='second'
		elif year_count ==3:
			year_count='third'
		elif year_count ==4:
			year_count='fourth'
		elif year_count ==5:
			year_count='fifth'
		elif year_count ==6:
			year_count='sixth'
		elif year_count ==7:
			year_count='seventh'
		elif year_count ==8:
			year_count='eighth'

		body_txt='''TO WHOM IT MAY CONCERN:\n\nThis is to certify that %s%s (passport No. %s) majoring in %s is studying in %s of Dali University since September, %s. %s has been promoted to %s %s academic year studying in the year of %s.\n\nSincerely yours\n'''%(Sex,Name,PPNo,Major,School,Grade,He,HeShe,year_count,toyear)
		infoNeed=list([Sex,Name,PPNo,Major,School,Grade,He,HeShe,year_count,toyear])
		if len([a for a in infoNeed if a.strip()==''])>0:
			print ('有缺失信息，请补全后再试！')
			showinfo('提示','有必填信息缺失，请补全后再试！')
		body=doc.add_paragraph()
		body_run=body.add_run(body_txt)
		paragraph_format = body.paragraph_format
		paragraph_format.line_spacing = docx.shared.Pt(30)
		inscription='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003, P. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8978   Fax:+86-872-221-8979'
		ins_run=body.add_run(inscription)
		font = ins_run.font
		font.bold = True

		today_appen = datetime.date.today()

		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
		else:
			doc.save('SC_EN_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')
			
	def loopNCC(self):#No Criminal Certificate
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[4],row[5]])
			if len([a for a in infoNeed if a.strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.NCC(doc,row,newPage)
		
	def loopMC(self):#Migration Certificate
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[4],row[5],row[6],row[7]])
			if len([a for a in infoNeed if a.strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.MC(doc,row,newPage)

		
	def loopSC_CN(self):#Study Certificate in Chinese language
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[3],row[4],row[5],row[7]])
			if len([a for a in infoNeed if a.strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.SC_CN(doc,row,newPage)

		
	def loopSC_EN(self):#Study Certificate in English language
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[4],row[5],row[6],row[7]])
			if len([a for a in infoNeed if a.strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.SC_EN(doc,row,newPage)
		
root=tk.Tk()
root.title('学生各类证明生成器')
#root.geometry('370x850')
#root.minsize(200, 200)

Select=tk.Label(root,text='\n请选择要生成的文件')
Select.pack()

app=Graduate(root)
root.mainloop()
