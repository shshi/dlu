#-*- coding: utf-8 -*-
import os
import time
import datetime
import docx
import xlrd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
import tkinter as tk
from tkinter.messagebox import *

class Graduate:
	def __init__(self,root):
		self.todolist=[]
		workbook = xlrd.open_workbook('CertInfo.xlsx')
		table = workbook.sheets()[0]
		nrows = table.nrows
		ncols = table.ncols
		for i in range(0,nrows):
			row_list=[]
			rowValues= table.row_values(i) 
			self.todolist.append(rowValues)
		self.todolist.remove(self.todolist[0])
		
		frame=tk.Frame(root)
		frame.pack()
		tk.Button(frame,text='No Criminal Certificate',command=self.loopNCC).pack(side=tk.LEFT)
		tk.Button(frame,text='Migration Certificate',command=self.loopMC).pack(side=tk.LEFT)
		tk.Button(frame,text='Study Certificate (Chinese)',command=self.loopSC_CN).pack(side=tk.LEFT)
		tk.Button(frame,text='Study Certificate (English)',command=self.loopSC_EN).pack(side=tk.LEFT)
		tk.Button(frame,text='Fee Structure',command=self.loopFS).pack(side=tk.LEFT)
		frame=tk.Frame(root)
		frame.pack()

		tk.Label(root,text='\nPowered by@ShaoTech\nCopyriht © 大理大学留学生教育服务中心 版权所有\n如有问题请联系：石少华，Email: shi.sh@foxmail.com').pack(side=tk.BOTTOM)
		README_btn=tk.Button(frame,text='使用说明',command=self.ReadMe)
		README_btn.pack(side=tk.TOP)

	def ReadMe(self):
		#os.system('README.txt')
		showinfo('使用说明','\n使用方法：\n1. 将待做学生信息写入CertInfo.xlsx\n2. 直接双击运行CertMaker.py\n\n*注意事项\n此程序所在文件夹中需有如下两个文件：\n1. 名为CertInfo.xlsx的Excel文件。此文件中为待做学生信息，直接用这里的模板即可。若无模板，请新建名为CertInfo.xlsx的Excel并设置第一行为表头，从左到右应依次为：\n序号，学号，护照姓名，中文名，性别(格式为Male/Female)，护照号，生日(英文)，国籍(英文)，银行名称，银行账号，银行代码。\n2. 名为template_blank.docx的大理大学留学生教育服务中心模版文件。若无此文件，生成的文件中将无含有大理大学信息的页眉页脚。\n')
		
	def NCC(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('CERTIFICATE OF NO CRIMINAL RECORD')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		date=doc.add_paragraph()
		date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		now = datetime.datetime.now()
		today=now.strftime("%B %d, %Y")
		date_txt='\n%s\n'%today
		#date_txt='\nJune 25'
		date.add_run(date_txt)
		#date.add_run('th').font.superscript = True
		#date.add_run(', 2019\n')
		
		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex or 'female' in Sex or '女' in Sex:
			Sex='Ms. '
			HeShe='her'
		else:
			Sex='Mr. '
			HeShe='his'
		Name=row[2].upper()
		try:
			PPNo='0'+str(int(row[5]))
		except:
			PPNo=row[5]	
		if row[1][0].upper()=='Y':
			Grade=row[1][1:5]
		elif row[1][2:4].upper()=='YY':
			Grade='20'+row[1][4:6]
		else:
			Grade='20'+row[1][2:4]
		body_txt='''This is to certify that %s%s (passport No. %s; student No. %s) has no disciplinary records against the rules and regulations of Dali University and has no Chinese judicial office records of committing any offense against Chinese criminal laws during %s study in P.R. China from September %s to till now.\n\nCertified by\n'''%(Sex,Name,PPNo,RollNo,HeShe,Grade)

		body=doc.add_paragraph()
		body_run=body.add_run(body_txt)
		body.paragraph_format.line_spacing = docx.shared.Pt(30)

		inscription_ZL='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8979   Fax:+86-872-221-8979'
		inscription_LM='\n\nProfessor Liu Ming\nDirector\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: mingliu192@aliyun.com\nTelephone: +86-872-221-8979   Fax: +86-872-221-8979'
		inscription_LFQ='\n\nLiu Fengqin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: dalijoe2005@aliyun.com\nTelephone: +86-872-221-8978   Fax: +86-872-221-8978'
		ins=doc.add_paragraph()
		ins.paragraph_format.line_spacing = docx.shared.Pt(25)
		ins_run=ins.add_run(inscription_LM)
		ins_run.font.bold = True

		today_appen = datetime.date.today()

		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
			
		else:
			doc.save('NCC_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')
			
	def MC(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('MIGRATION CERTIFICATE')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		date=doc.add_paragraph()
		date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		now = datetime.datetime.now()
		today=now.strftime("%B %d, %Y")
		date_txt='\n%s\n'%today
		#date_txt='\nJune 25'
		date.add_run(date_txt)
		#date.add_run('th').font.superscript = True
		#date.add_run(', 2019\n')

		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex or 'female' in Sex or '女' in Sex:
			Sex='Ms. '
			HeShe='her'
		else:
			Sex='Mr. '
			HeShe='his'
		Name=row[2].upper()
		try:
			PPNo='0'+str(int(row[5]))
		except:
			PPNo=row[5]
		DOB=row[6]
		Nation=row[7]

		body_txt='''This university has no objection to the admission of %s%s, from %s, bearing Dali University’s registration No. %s, passport No. %s, born on %s, for %s further study in any institution or university in any country.\n\nWe wish %s success in life.\n\n\n'''%(Sex,Name,Nation,RollNo,PPNo,DOB,HeShe,HeShe)
		body=doc.add_paragraph()
		body_run=body.add_run(body_txt)
		paragraph_format = body.paragraph_format
		paragraph_format.line_spacing = docx.shared.Pt(30)
		inscription_ZL='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8979   Fax:+86-872-221-8979'
		inscription_LM='\n\nProfessor Liu Ming\nDirector\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: mingliu192@aliyun.com\nTelephone: +86-872-221-8979   Fax: +86-872-221-8979'
		inscription_LFQ='\n\nLiu Fengqin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: dalijoe2005@aliyun.com\nTelephone: +86-872-221-8978   Fax: +86-872-221-8978'
		ins_run=body.add_run(inscription_LM)
		font = ins_run.font
		font.bold = True

		today_appen = datetime.date.today()

		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
			
		else:
			doc.save('MC_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')

	def SC_CN(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('在 读 证 明\n\n')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		
		RollNo=row[1]
		if row[1][0].upper()=='Y':
			Grade=row[1][1:5]
		elif row[1][2:4].upper()=='YY':
			Grade='20'+row[1][4:6]
		else:
			Grade='20'+row[1][2:4]
		Sex=row[4]
		if 'Female' in Sex or 'female' in Sex or '女' in Sex:
			Sex='女'
		else:
			Sex='男'
		Name=row[2].upper()
		Name_CN=row[3]
		try:
			PPNo='0'+str(int(row[5]))
		except:
			PPNo=row[5]
		Nation=row[7]
		if 'India' in Nation or '印度' in Nation:
			Nation='印度'
		elif 'Nepal' in Nation or '尼泊尔' in Nation:
			Nation='尼泊尔'
		elif 'Pakistan' in Nation or '巴基斯坦' in Nation:
			Nation='巴基斯坦'
		elif 'Bangladesh' in Nation or '孟加拉国' in Nation:
			Nation='孟加拉国'
		elif 'Ivory Cost' in Nation or '科特迪瓦' in Nation:
			Nation='科特迪瓦'
		elif 'Laos' in Nation or '老挝' in Nation:
			Nation='老挝'
		elif 'Cambodia' in Nation or '柬埔寨' in Nation:
			Nation='柬埔寨'
		elif 'Tanzania' in Nation or '坦桑尼亚' in Nation:
			Nation='坦桑尼亚'
		elif 'Viet' in Nation or '越南' in Nation:
			Nation='越南'
		elif 'Somalia' in Nation or '索马里' in Nation:
			Nation='索马里'
		elif 'Burma' in Nation or 'Myanmar' in Nation or '缅甸' in Nation:
			Nation='缅甸'
		elif 'Zambia' in Nation or '赞比亚' in Nation:
			Nation='赞比亚'
		elif 'Yemen' in Nation or '也门' in Nation:
			Nation='也门'
		elif 'Mongolia' in Nation or '蒙古' in Nation:
			Nation='蒙古'

		body=doc.add_paragraph()
		body.paragraph_format.line_spacing = docx.shared.Pt(30)
		
		body.add_run('    兹证明')
		body.add_run('%s'%Name).font.underline = True

		body.add_run('，性别')
		body.add_run('%s'%Sex).font.underline = True
		
		body.add_run('，中文名')
		body.add_run('%s'%Name_CN).font.underline = True

		body.add_run('，国籍')
		body.add_run('%s'%Nation).font.underline = True

		body.add_run('，护照号码')
		body.add_run('%s'%PPNo).font.underline = True

		body.add_run('，为我校')
		body.add_run('临床医学院%s级临床医学专业本科生'%Grade).font.underline = True
		body.add_run('。该生于%s年10月入学，%s年7月从我校毕业。\n\n\t特此证明。\n\n'%(Grade,str(int(Grade)+6)))

		year='年'
		month='月'
		day='日'
		cc=time.localtime(time.time())
		end_txt='大理大学留学生教育服务中心\n%s'%str(cc.tm_year)+year+str(cc.tm_mon)+month+str(cc.tm_mday)+day
		#end_txt='大理大学留学生教育服务中心\n2019年6月25日'
		end=doc.add_paragraph()
		end_run=end.add_run(end_txt)
		paragraph_format = end.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		today_appen = datetime.date.today()
		
		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
			
		else:
			doc.save('SC_CN_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')

	def SC_EN(self,doc,row,newPage):
		title=doc.add_paragraph()
		title_run=title.add_run('STUDY CERTIFICATE')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		date=doc.add_paragraph()
		date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		now = datetime.datetime.now()
		today=now.strftime("%B %d, %Y")
		date_txt='\n%s\n'%today
		#date_txt='\nJune 25'
		date.add_run(date_txt)
		#date.add_run('th').font.superscript = True
		#date.add_run(', 2019\n')

		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex or 'female' in Sex or '女' in Sex:
			Sex='Ms. '
			HeShe='her'
			He='She'
		else:
			Sex='Mr. '
			HeShe='his'
			He='He'
		Name=row[2].upper()
		try:
			PPNo='0'+str(int(row[5]))
		except:
			PPNo=row[5]
		Nation=row[7]
		Major='Clinical Medicine'
		School='the College of Clinical Medicine'
		if row[1][0].upper()=='Y':
			Grade=row[1][1:5]
		elif row[1][2:4].upper()=='YY':
			Grade='20'+row[1][4:6]
		else:
			Grade='20'+row[1][2:4]
		toyear=now.strftime("%Y")
		year_count=int(toyear)-int(Grade)
		if year_count ==0:
			year_count='first'
		elif year_count ==1:
			year_count='first'
		elif year_count ==2:
			year_count='second'
		elif year_count ==3:
			year_count='third'
		elif year_count ==4:
			year_count='fourth'
		elif year_count ==5:
			year_count='fifth'
		elif year_count ==6:
			year_count='sixth'
		elif year_count ==7:
			year_count='seventh'
		elif year_count ==8:
			year_count='eighth'

		body_txt='''TO WHOM IT MAY CONCERN:\n\nThis is to certify that %s%s (passport No. %s), majoring in %s, has been studying in %s, Dali University since September, %s. %s has been promoted to %s %s academic year in the year of %s.\n\nSincerely yours,\n'''%(Sex,Name,PPNo,Major,School,Grade,He,HeShe,year_count,toyear)
		body_txt_OneYear=''
		body=doc.add_paragraph()
		body_run=body.add_run(body_txt)
		paragraph_format = body.paragraph_format
		paragraph_format.line_spacing = docx.shared.Pt(30)
		inscription_ZL='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8979   Fax:+86-872-221-8979'
		inscription_LM='\n\nProfessor Liu Ming\nDirector\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: mingliu192@aliyun.com\nTelephone: +86-872-221-8979   Fax: +86-872-221-8979'
		inscription_LFQ='\n\nLiu Fengqin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: dalijoe2005@aliyun.com\nTelephone: +86-872-221-8978   Fax: +86-872-221-8978'
		ins=doc.add_paragraph()
		ins.paragraph_format.line_spacing = docx.shared.Pt(25)
		ins_run=ins.add_run(inscription_LM)
		ins_run.font.bold = True

		today_appen = datetime.date.today()

		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
		else:
			doc.save('SC_EN_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')

	def FS(self,doc,row,newPage):
		title=doc.add_paragraph()
		#title.space_after = docx.shared.Pt(0)
		#title.space_before = docx.shared.Pt(0)
		title_run=title.add_run('To Whom It May Concern')
		font = title_run.font
		#font.name = 'Calibri'
		font.bold = True
		font.size = docx.shared.Pt(22)
		paragraph_format = title.paragraph_format
		paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		paragraph_format.space_before = docx.shared.Pt(0)

		date=doc.add_paragraph()
		date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		now = datetime.datetime.now()
		today=now.strftime("%B %d, %Y")
		date_txt='%s'%today
		#date_txt='\nJune 25'
		date.add_run(date_txt).font.size = docx.shared.Pt(13)
		#date.add_run('th').font.superscript = True
		#date.add_run(', 2019\n')
		
		RollNo=row[1]
		Sex=row[4]
		if 'Female' in Sex or 'female' in Sex or '女' in Sex:
			Sex='Ms. '
			HeShe='her'
		else:
			Sex='Mr. '
			HeShe='his'
		Name=row[2].upper()
		try:
			PPNo='0'+str(int(row[5]))
		except:
			PPNo=row[5]	
		if row[1][0].upper()=='Y':
			Grade=row[1][1:5]
			PG='a post graduate'
		elif row[1][2:4].upper()=='YY':
			Grade='20'+row[1][4:6]
			PG='an undergraduate'
		else:
			Grade='20'+row[1][2:4]
			PG='an undergraduate'
			
		toyear=now.strftime("%Y")
		year_count=int(toyear)-int(Grade)
		Bank=row[8]
		BanckAcc=row[9]
		SwiftCode=row[10]
		body_txt='''This is to certify that %s%s (passport No. %s) is %s student of %s batch of Dali University. The fee structure for academic year 2018–2019 is as per below. The exchange rate of US Dollar into Chinese Yuan for this academic year is 6.41 (according to the exchange rate provided by Bank of China on May 31th, 2018).'''%(Sex,Name,PPNo,PG,Grade)

		body=doc.add_paragraph()
		body_run=body.add_run(body_txt).font.size = docx.shared.Pt(13)
		body.paragraph_format.line_spacing = docx.shared.Pt(20)
		paragraph_format=body.paragraph_format
		paragraph_format.space_after = docx.shared.Pt(10)
		
		if PG=='an undergraduate':
			if year_count>1:
				if int(Grade)<2011:
					Tuition='13600RMB'
					AccmdFee='2000RMB'
					RegFee='240RMB'
				if int(Grade)==2011:
					Tuition='16000RMB'
					AccmdFee='2000RMB'
					RegFee='240RMB'
				if int(Grade)==2012:
					Tuition='17000RMB'
					AccmdFee='2000RMB'
					RegFee='240RMB'
				if int(Grade)==2013:
					Tuition='18000RMB'
					AccmdFee='2000RMB'
					RegFee='240RMB'
				if int(Grade)==2014:
					Tuition='23000RMB'
					AccmdFee='2000RMB'
					RegFee='300RMB'
				if int(Grade)==2015:
					Tuition='16000RMB'
					AccmdFee='2200RMB'
					RegFee='300RMB'
				if int(Grade)==2016 or int(Grade)==2017:
					Tuition='17000RMB'
					AccmdFee='2200RMB'
					RegFee='300RMB'
				if int(Grade)==2018:
					Tuition='18000RMB'
					AccmdFee='2200RMB'
					RegFee='300RMB'
				if int(Grade)==2019:
					Tuition='20000RMB'
					AccmdFee='2200RMB'
					RegFee='300RMB'
			else:
				if int(Grade)==2018:
					Tuition='18000RMB'
					AccmdFee='2200RMB'
					RegFee='3000RMB'
				if int(Grade)==2019:
					Tuition='20000RMB'
					AccmdFee='2200RMB'
					RegFee='3000RMB'
		else:
			if year_count>1:
				if int(Grade)==2012 or int(Grade)==2013:
					Tuition='22000RMB'
					AccmdFee='2400RMB'
					RegFee='240RMB'
				else:
					Tuition='28000RMB'
					AccmdFee='2400RMB'
					RegFee='240RMB'
			else:
				Tuition='28000RMB'
				AccmdFee='2400RMB'
				RegFee='2400RMB'
				
		DineFee,PocketFee,TranspFee,InsFee="3450RMB*","16500RMB*","6850RMB*","600RMB"
		TotalAmount=str(int(RegFee.replace('RMB',''))+int(Tuition.replace('RMB',''))+int(AccmdFee.replace('RMB',''))+int(DineFee.replace('RMB*',''))+int(PocketFee.replace('RMB*',''))+int(TranspFee.replace('RMB*',''))+int(InsFee.replace('RMB','')))+'RMB'
		SelfAmount=str(int(DineFee.replace('RMB*',''))+int(PocketFee.replace('RMB*',''))+int(TranspFee.replace('RMB*','')))+'RMB'
		SchoolAmount=str(int(RegFee.replace('RMB',''))+int(Tuition.replace('RMB',''))+int(AccmdFee.replace('RMB',''))+int(InsFee.replace('RMB','')))+'RMB'
		
		table = doc.add_table(rows=9,cols=3,style='Table Grid')
		table.alignment=WD_TABLE_ALIGNMENT.CENTER
		#table.style.font.size=docx.shared.Pt(12)
		#table.autofit = True
		for cell in table.columns[0].cells:
			cell.width = docx.shared.Inches(1)
			cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		for cell2 in table.columns[2].cells:
			cell2.width = docx.shared.Inches(2.5)

		table.cell(0,0).paragraphs[0].add_run('Serial No.').font.bold = True
		table.cell(0,1).paragraphs[0].add_run('Contents').font.bold = True
		table.cell(0,2).paragraphs[0].add_run('Amount').font.bold = True

		
		table.cell(1,0).text = "1"
		table.cell(1,1).text = "Registration Fee"
		table.cell(1,2).text = RegFee
		
		table.cell(2,0).text = "2"
		table.cell(2,1).text = "Medical Tuition Fee"
		table.cell(2,2).text = Tuition
		
		table.cell(3,0).text = "3"
		table.cell(3,1).text = "Accommodation Fee"
		if PG=='an undergraduate':
			table.cell(3,2).text = AccmdFee
		else:
			table.cell(3,2).text = AccmdFee+' (if live in the campus)'
		
		table.cell(4,0).text = "4"
		table.cell(4,1).text = "Dining Fee"
		table.cell(4,2).text = DineFee
		
		table.cell(5,0).text = "5"
		table.cell(5,1).text = "Pocket Money"
		table.cell(5,2).text = PocketFee
		
		table.cell(6,0).text = "6"
		table.cell(6,1).text = "Airfare & Transportation"
		table.cell(6,2).text = TranspFee
		
		table.cell(7,0).text = "7"
		table.cell(7,1).text = "Insurance"
		table.cell(7,2).text = InsFee
		
		table.cell(8,0).text = ""
		table.cell(8,1).text = "Total Amount"
		table.cell(8,2).text = TotalAmount

		for i in range(9):
			for cell in table.rows[i].cells:
				cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
				cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		for row in table.rows:
			row.height = docx.shared.Inches(0.25)
		
		Note_Bank='''Those items with asterisks are suggested amount of fees for this academic year. Dali University has no objection to transferring all these fees to %s personal account.\n\nAccount Name: %s\nDeposit Bank: %s\nBank Account No: %s\nSwift Code: %s\n\n'''%(HeShe,Name,Bank,BanckAcc,SwiftCode)
		
		NB=doc.add_paragraph()
		NB.paragraph_format.line_spacing = docx.shared.Pt(17)
		NB.add_run('Note: ').font.bold=True
		NB.add_run(Note_Bank).font.size=docx.shared.Pt(11)

		inscription_ZL='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students, Dali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8979   Fax:+86-872-221-8979'
		inscription_LM='\n\nProfessor Liu Ming\nDirector\nEducation & Service Center for International Students, Dali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: mingliu192@aliyun.com\nTelephone: +86-872-221-8979   Fax: +86-872-221-8979'
		inscription_LFQ='\n\nLiu Fengqin\nDeputy Director\nEducation & Service Center for International Students, Dali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003\nP. R. CHINA\nE-mail: dalijoe2005@aliyun.com\nTelephone: +86-872-221-8978   Fax: +86-872-221-8978'
		ins=doc.add_paragraph()
		ins.paragraph_format.line_spacing = docx.shared.Pt(20)
		ins_run=ins.add_run(inscription_LM)
		ins_run.font.bold = True

		today_appen = datetime.date.today()

		if newPage=='Y':
			doc.add_page_break()
			doc.add_paragraph()
			
		else:
			doc.save('FS_%s.docx'%today_appen)
			showinfo('提示','生成完毕！')

	def log(self,item):
		list=self.todolist
		logtime=datetime.datetime.now()
		#logtime=time.strftime('%Y.%m.%d %H:%M:%S ',time.localtime(time.time()))
		with open('log.log','a') as p:
			for data in list:
				p.write(data[1]+'\t'+data[2]+'\t'+item+'\t'+str(logtime)+'\n')

	def loopNCC(self):#No Criminal Certificate
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[4],row[5]])
			if len([a for a in infoNeed if str(a).strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.NCC(doc,row,newPage)
		self.log('NCC')
		
	def loopMC(self):#Migration Certificate
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[4],row[5],row[6],row[7]])
			if len([a for a in infoNeed if str(a).strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.MC(doc,row,newPage)
		self.log('MC')
		
	def loopSC_CN(self):#Study Certificate in Chinese language
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[3],row[4],row[5],row[7]])
			if len([a for a in infoNeed if str(a).strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.SC_CN(doc,row,newPage)
		self.log('SC_CN')
		
	def loopSC_EN(self):#Study Certificate in English language
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(16)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[4],row[5],row[7]])
			if len([a for a in infoNeed if str(a).strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.SC_EN(doc,row,newPage)
		self.log('SC_EN')

	def loopFS(self):#Study Certificate in English language
		doc = Document('.\\template_blank.docx')
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = docx.shared.Pt(12)
		for row in self.todolist:
			infoNeed=list([row[1],row[2],row[4],row[5],row[8],row[9],row[10]])
			if len([a for a in infoNeed if str(a).strip()==''])>0:
				showinfo('提示','有必填信息缺失，请补全后再试！')
				break
			if row!=self.todolist[-1]:
				newPage='Y'
			else:
				newPage='N'
			self.FS(doc,row,newPage)
		self.log('FS')

		
root=tk.Tk()
root.title('学生各类证明生成器')
#root.geometry('370x850')
#root.minsize(200, 200)

Select=tk.Label(root,text='\n请选择要生成的文件')
Select.pack()

app=Graduate(root)
root.mainloop()