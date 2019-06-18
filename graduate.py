#-*- coding: utf-8 -*-
import time
import datetime
import docx
import xlrd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def todo():
	todolist=[]
	workbook = xlrd.open_workbook('graduate.xlsx')
	table = workbook.sheets()[0]
	nrows = table.nrows
	ncols = table.ncols
	for i in range(0,nrows):
		row_list=[]
		rowValues= table.row_values(i) 
		todolist.append(rowValues)
	todolist.remove(todolist[0])
	#print (todolist)
	return todolist
	
def NCC(doc,row,newPage):
	title=doc.add_paragraph()
	title_run=title.add_run('CERTIFICATE OF NO CRIMINAL RECORD')
	font = title_run.font
	#font.name = 'Calibri'
	font.bold = True
	font.size = docx.shared.Pt(22)
	paragraph_format = title.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	date=doc.add_paragraph()
	date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	now = datetime.datetime.now()
	today=now.strftime("%B %d, %Y")
	date_txt='\n%s\n'%today
	date_txt='\nJune 25'
	date.add_run(date_txt)
	date.add_run('th').font.superscript = True
	date.add_run(', 2019\n')
	
	RollNo=row[1]
	Sex=row[4]
	if 'Female' in Sex:
		Sex='Ms. '
		HeShe='her'
	else:
		Sex='Mr. '
		HeShe='his'
	Name=row[2]
	PPNo=row[5]
	Grade='20'+row[1][2:4]
	if row[1][0]=='Y' or row[1][0]=='y':
		Grade=row[1][1:5]
	body_txt='''This is to certify that %s%s (passport No. %s; student No. %s) has no disciplinary records against the rules and regulations of Dali University and has no Chinese judicial office records of committing any offense against Chinese criminal laws during %s study in P.R. China from September %s to till now.\n\nCertified by\n'''%(Sex,Name,PPNo,RollNo,HeShe,Grade)
	body=doc.add_paragraph()
	body_run=body.add_run(body_txt)
	body.paragraph_format.line_spacing = docx.shared.Pt(30)

	inscription='\n\n\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003, P. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8978   Fax:+86-872-221-8979'	
	ins=doc.add_paragraph()
	ins.paragraph_format.line_spacing = docx.shared.Pt(25)
	ins_run=ins.add_run(inscription)
	ins_run.font.bold = True

	today_appen = datetime.date.today()

	if newPage=='Y':
		doc.add_page_break()
		doc.add_paragraph()
		
	else:
		doc.save('NCC_%s.docx'%today_appen)
		
def MC(doc,row,newPage):
	title=doc.add_paragraph()
	title_run=title.add_run('MIGRATION CERTIFICATE')
	font = title_run.font
	#font.name = 'Calibri'
	font.bold = True
	font.size = docx.shared.Pt(22)
	paragraph_format = title.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	date=doc.add_paragraph()
	date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	now = datetime.datetime.now()
	today=now.strftime("%B %d, %Y")
	date_txt='\n%s\n'%today
	date_txt='\nJune 25'
	date.add_run(date_txt)
	date.add_run('th').font.superscript = True
	date.add_run(', 2019\n')

	RollNo=row[1]
	Sex=row[4]
	if 'Female' in Sex:
		Sex='Ms. '
		HeShe='her'
	else:
		Sex='Mr. '
		HeShe='his'
	Name=row[2]
	PPNo=row[5]
	DOB=row[6]
	Nation=row[7]

	body_txt='''This university has no objection to the admission of %s%s, from %s, bearing Dali University’s registration No. %s, passport No. %s, born on %s, for %s further study in any institution or university in any country.\n\nWe wish %s success in life.\n\n\n'''%(Sex,Name,Nation,RollNo,PPNo,DOB,HeShe,HeShe)
	body=doc.add_paragraph()
	body_run=body.add_run(body_txt)
	paragraph_format = body.paragraph_format
	paragraph_format.line_spacing = docx.shared.Pt(30)
	inscription='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003, P. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8978   Fax:+86-872-221-8979'
	ins_run=body.add_run(inscription)
	font = ins_run.font
	font.bold = True

	today_appen = datetime.date.today()

	if newPage=='Y':
		doc.add_page_break()
		doc.add_paragraph()
		
	else:
		doc.save('MC_%s.docx'%today_appen)

def SC_CN(doc,row,newPage):
	title=doc.add_paragraph()
	title_run=title.add_run('在 读 证 明\n\n')
	font = title_run.font
	#font.name = 'Calibri'
	font.bold = True
	font.size = docx.shared.Pt(22)
	paragraph_format = title.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	RollNo=row[1]
	Sex=row[4]
	if 'Female' in Sex:
		Sex='女'
	else:
		Sex='男'
	Name=row[2]
	Name_CN=row[3]
	PPNo=row[5]
	Nation=row[7]
	if 'India' in Nation:
		Nation='印度'
	elif 'Nepal' in Nation:
		Nation='尼泊尔'
	elif 'Pakistan' in Nation:
		Nation='巴基斯坦'
	elif 'Bangladesh' in Nation:
		Nation='孟加拉国'
	elif 'Ivory Cost' in Nation:
		Nation='科特迪瓦'

	body=doc.add_paragraph()
	body.paragraph_format.line_spacing = docx.shared.Pt(30)
	
	body.add_run('    兹证明')
	body.add_run('%s'%Name).font.underline = True

	body.add_run('，性别')
	body.add_run('%s'%Sex).font.underline = True
	
	body.add_run('，中文名')
	body.add_run('%s'%Name_CN).font.underline = True

	body.add_run('，国籍')
	body.add_run('%s'%Nation).font.underline = True

	body.add_run('，护照号码')
	body.add_run('%s'%PPNo).font.underline = True

	body.add_run('，为我校')
	body.add_run('临床医学院2013级临床医学专业本科生').font.underline = True
	body.add_run('。该生于2013年10月入学，2019年7月从我校毕业。\n\n\t特此证明。\n\n')

	year='年'
	month='月'
	day='日'
	cc=time.localtime(time.time())
	end_txt='大理大学留学生教育服务中心\n%s'%str(cc.tm_year)+year+str(cc.tm_mon)+month+str(cc.tm_mday)+day
	#end_txt='大理大学留学生教育服务中心\n2019年6月25日'
	end=doc.add_paragraph()
	end_run=end.add_run(end_txt)
	paragraph_format = end.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	today_appen = datetime.date.today()
	
	if newPage=='Y':
		doc.add_page_break()
		doc.add_paragraph()
		
	else:
		doc.save('SC_CN_%s.docx'%today_appen)

def SC_EN(doc,row,newPage):
	title=doc.add_paragraph()
	title_run=title.add_run('STUDY CERTIFICATE')
	font = title_run.font
	#font.name = 'Calibri'
	font.bold = True
	font.size = docx.shared.Pt(22)
	paragraph_format = title.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	date=doc.add_paragraph()
	date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	now = datetime.datetime.now()
	today=now.strftime("%B %d, %Y")
	date_txt='\n%s\n'%today
	#date_txt='\nJune 25'
	date.add_run(date_txt)
	#date.add_run('th').font.superscript = True
	#date.add_run(', 2019\n')

	RollNo=row[1]
	Sex=row[4]
	if 'Female' in Sex:
		Sex='Ms. '
		HeShe='her'
		He='She'
	else:
		Sex='Mr. '
		HeShe='his'
		He='He'
	Name=row[2]
	PPNo=row[5]
	DOB=row[6]
	Nation=row[7]
	Major='Clinical Medicine'
	School='Clinical Medicine College'
	Grade='20'+row[1][2:4]
	if row[1][0]=='Y' or row[1][0]=='y':
		Grade=row[1][1:5]
	toyear=now.strftime("%Y")
	year_count=int(toyear)-int(Grade)
	if year_count ==0:
		year_count='first'
	elif year_count ==1:
		year_count='first'
	elif year_count ==2:
		year_count='second'
	elif year_count ==3:
		year_count='third'
	elif year_count ==4:
		year_count='fourth'
	elif year_count ==5:
		year_count='fifth'
	elif year_count ==6:
		year_count='sixth'
	elif year_count ==7:
		year_count='seventh'
	elif year_count ==8:
		year_count='eighth'

	body_txt='''TO WHOM IT MAY CONCERN:\n\nThis is to certify that %s%s (passport No. %s) majoring in %s is studying in %s of Dali University since September, %s. %s has been promoted to %s %s academic year studying in the year of %s.\n\nSincerely yours\n'''%(Sex,Name,PPNo,Major,School,Grade,He,HeShe,year_count,toyear)
	body=doc.add_paragraph()
	body_run=body.add_run(body_txt)
	paragraph_format = body.paragraph_format
	paragraph_format.line_spacing = docx.shared.Pt(30)
	inscription='\n\nMs. Zhou Lin\nDeputy Director\nEducation & Service Center for International Students\nDali University\nNo. 2 Hongsheng Road, Dali, Yunnan 671003, P. R. CHINA\nEmail: leanne927cn@hotmail.com\nTelephone: +86-872-221-8978   Fax:+86-872-221-8979'
	ins_run=body.add_run(inscription)
	font = ins_run.font
	font.bold = True

	today_appen = datetime.date.today()

	if newPage=='Y':
		doc.add_page_break()
		doc.add_paragraph()
		
	else:
		doc.save('SC_EN_%s.docx'%today_appen)

		
def loopNCC():#No Criminal Certificate
	todolist=todo()
	doc = Document('.\\template_blank.docx')
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = docx.shared.Pt(16)
	for row in todolist:
		if row!=todolist[-1]:
			newPage='Y'
		else:
			newPage='N'
		NCC(doc,row,newPage)

def loopMC():#Migration Certificate
	todolist=todo()
	doc = Document('.\\template_blank.docx')
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = docx.shared.Pt(16)
	for row in todolist:
		if row!=todolist[-1]:
			newPage='Y'
		else:
			newPage='N'
		MC(doc,row,newPage)

def loopSC_CN():#Study Certificate in Chinese language
	todolist=todo()
	doc = Document('.\\template_blank.docx')
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = docx.shared.Pt(16)
	for row in todolist:
		if row!=todolist[-1]:
			newPage='Y'
		else:
			newPage='N'
		SC_CN(doc,row,newPage)

def loopSC_EN():#Study Certificate in English language
	todolist=todo()
	doc = Document('.\\template_blank.docx')
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = docx.shared.Pt(16)
	for row in todolist:
		if row!=todolist[-1]:
			newPage='Y'
		else:
			newPage='N'
		SC_EN(doc,row,newPage)

loopNCC()
loopMC()
loopSC_CN()
loopSC_EN()
print ('\ndone')
time.sleep(3)