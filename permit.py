#-*- coding: utf-8 -*-
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlrd
import time
import datetime
import re
todolist=[]
workbook = xlrd.open_workbook('todolist.xlsx')
table = workbook.sheets()[0]
nrows = table.nrows
ncols = table.ncols
new=input('此批学生是否为新生？（请输入Y或者N）：')
while new != "Y" and new != "N":
	new = input('输入错误，请请输入Y或者N（大写）：')

for i in range(0,nrows):
	row_list=[]
	rowValues= table.row_values(i) 
	todolist.append(rowValues)
todolist.remove(todolist[0])
#print (todolist)

doc = Document('.\\template_blank.docx')
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = docx.shared.Pt(16)
	
for row in todolist:
	print (row)
	ymd_permit=str(row[7]).replace('/','.')
	ymd_permit=ymd_permit.split('.')
	if int(ymd_permit[1])==1:
		ymd_todo=ymd_permit[0]+'年'+'12'+'月'+ymd_permit[-1]+'日'
		ymd_todo_bijiao=ymd_permit[0]+'.12'
		ymd_todo_bijiao=ymd_todo_bijiao.split('.')
	else:
		ymd_todo=str(int(ymd_permit[0])+1)+'年'+str(int(ymd_permit[1])-1)+'月'+ymd_permit[-1]+'日'
		ymd_todo_bijiao=str(int(ymd_permit[0])+1)+'.'+str(int(ymd_permit[1])-1)
		ymd_todo_bijiao=ymd_todo_bijiao.split('.')
	ymd_permit=ymd_permit[0]+'年'+ymd_permit[1]+'月'+ymd_permit[-1]+'日'
	#print (ymd_todo)

	title=doc.add_paragraph()
	title_run=title.add_run('证  明\n\n')
	font = title_run.font
	#font.name = 'Calibri'
	font.bold = True
	font.size = docx.shared.Pt(22)
	paragraph_format = title.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	date_txt='大理州公安局出入境管理支队:'
	date=doc.add_paragraph()
	date_run=date.add_run(date_txt)
	font = date_run.font
	font.size = docx.shared.Pt(18)
	#paragraph_format = date.paragraph_format
	#paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	if str(row[8])=='':
		JW202='X年X月X日'
	else:
		ymd_202=str(row[8]).replace('/','.')
		ymd_202_bijiao=ymd_202
		ymd_202_bijiao=ymd_202_bijiao.split('.')
		ymd_202=ymd_202.split('.')
		ymd_202=ymd_202[0]+'年'+ymd_202[1]+'月'
		JW202=ymd_202
		#print (ymd_202_bijiao[0],ymd_todo_bijiao[0])
		if ymd_202_bijiao[0]==ymd_todo_bijiao[0] and int(ymd_202_bijiao[1])<int(ymd_todo_bijiao[1]):
			if int(ymd_202_bijiao[1]) in [1,3,5,7,8,10,12]:
				ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'31日'
			elif int(ymd_202_bijiao[1])==2:
				if (int(ymd_202_bijiao[0])%4==0 and int(ymd_202_bijiao[0])%100!=0) or (int(ymd_202_bijiao[0])%400==0):
					ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'29日'
				else:
					ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'28日'
			else:
				ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'30日'
			with open('log.txt','a') as p:			
				p.write('JW202需更新：'+row[1]+'\n')
		elif ymd_202_bijiao[0]<ymd_todo_bijiao[0]:
			if ymd_202_bijiao[1] in [1,3,5,7,8,10,12]:
				ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'31日'
			elif ymd_202_bijiao[1]==2:
				if (int(ymd_202_bijiao[0])%4==0 and int(ymd_202_bijiao[0])%100!=0) or (int(ymd_202_bijiao[0])%400==0):
					ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'29日'
				else:
					ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'28日'
			else:
				ymd_todo=ymd_202_bijiao[0]+'年'+ymd_202_bijiao[1]+'月'+'30日'	
			with open('log.txt','w') as p:
				p.write('JW202需更新：'+row[1]+'\n')

	try:
		ppNo='0'+str(int(row[6]))
	except:
		ppNo=row[6]
	


	if new=='Y':
		'''body_txt="    兹有我校 %s 籍学生 %s，性别: %s, 护照号码为 %s。其居留许可到期，现需办理学习居留许可延期手续，时间为%s，请按有关规定给予办理为谢。\n    JW202表到期时间：%s\n\n\n\n"%(row[5],row[2].upper(),row[4],ppNo,ymd_todo,JW202)
		body=doc.add_paragraph()
		body_run=body.add_run(body_txt)
		paragraph_format = body.paragraph_format
		paragraph_format.line_spacing = docx.shared.Pt(30)
		font.underline = False'''#不加下划线
		body=doc.add_paragraph()
		
		body.add_run('    兹有我校')
		paragraph_format = body.paragraph_format
		paragraph_format.line_spacing = docx.shared.Pt(30)
		body.add_run(' %s '%row[5]).font.underline = True
		
		body.add_run('籍学生')
		body.add_run('%s'%row[2].upper()).font.underline = True

		body.add_run(', 性别：')
		body.add_run('%s'%row[4]).font.underline = True

		body.add_run('，护照号码为')
		body.add_run('%s'%ppNo).font.underline = True

		body.add_run('。其居留许可到期，现需办理学习居留许可延期手续，时间为')
		body.add_run('%s'%ymd_todo).font.underline = True

		body.add_run('，请按有关规定给予办理为谢。\n    居留许可到期时间：')
		body.add_run('%s'%ymd_permit).font.underline = True

		body.add_run('\n\n\n\n')
	else:
		body=doc.add_paragraph()
		
		body.add_run('    兹有我校')
		paragraph_format = body.paragraph_format
		paragraph_format.line_spacing = docx.shared.Pt(30)
		body.add_run(' %s '%row[5]).font.underline = True
		
		body.add_run('籍学生')
		body.add_run('%s'%row[2].upper()).font.underline = True
		
		body.add_run(', 性别：')
		body.add_run('%s'%row[4]).font.underline = True
		if row[9]=='':
			body.add_run('，护照号码为')
			body.add_run('%s'%ppNo).font.underline = True
		else:
			body.add_run('，旧护照号码为')
			body.add_run('%s'%ppNo).font.underline = True
			body.add_run('，新护照号码为')
			body.add_run('%s'%row[9]).font.underline = True
		
		body.add_run('。其居留许可到期，现需办理学习居留许可延期手续，时间为')
		body.add_run('%s'%ymd_todo).font.underline = True

		body.add_run('，请按有关规定给予办理为谢。\n    居留许可到期时间：')
		body.add_run('%s'%ymd_permit).font.underline = True

		body.add_run('\n    JW202表到期时间：')
		body.add_run('%s'%JW202).font.underline = True

		body.add_run('\n\n\n\n')

	year='年'
	month='月'
	day='日'
	cc=time.localtime(time.time())
	end_txt='大理大学留学生教育服务中心\n%s'%str(cc.tm_year)+year+str(cc.tm_mon)+month+str(cc.tm_mday)+day
	end=doc.add_paragraph()
	end_run=end.add_run(end_txt)
	paragraph_format = end.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	if row!=todolist[-1]:
		doc.add_page_break()
		doc.add_paragraph()

today = datetime.date.today()
#today=str(today).replace('-','')
doc.save('permit_%s.docx'%today)
print ('已成功生成文件！')
time.sleep(7)
